import argparse
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=100, right=100):
    """Sets margins (padding) inside a table cell in DXA units."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    """Fills cell background."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_jip_template(filename="template_jip_filled.docx"):
    """Generates a document conforming to the JIP (Jurnal Informatika Polinema) styling guidelines."""
    try:
        doc = Document()
        
        # JIP Page Margins: Top 30mm (1.18 in), Left 30mm (1.18 in), Bottom 20mm (0.78 in), Right 20mm (0.78 in)
        section = doc.sections[0]
        section.top_margin = Inches(1.18)
        section.left_margin = Inches(1.18)
        section.bottom_margin = Inches(0.78)
        section.right_margin = Inches(0.78)
        
        # Base Style settings
        style_normal = doc.styles['Normal']
        style_normal.font.name = 'Times New Roman'
        style_normal.font.size = Pt(10)
        style_normal.font.color.rgb = RGBColor(0, 0, 0)
        
        # --- Section 1: Title and Metadata (Single Column) ---
        
        # JIP Journal ISSN Header line (Simulated)
        issn_p = doc.add_paragraph()
        issn_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        issn_run = issn_p.add_run("JIP (Jurnal Informatika Polinema)  ISSN: 2614-6371 E-ISSN: 2407-070X")
        issn_run.font.name = 'Times New Roman'
        issn_run.font.size = Pt(9)
        issn_run.font.italic = True
        
        # Title (14pt, Bold, Centered)
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(24)
        title_p.paragraph_format.space_after = Pt(12)
        title_run = title_p.add_run("Judul Makalah Jurnal Informatika Polinema\n(Maksimum 10 Kata)")
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        
        # Authors (10pt, Bold, Centered)
        auth_p = doc.add_paragraph()
        auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        auth_p.paragraph_format.space_after = Pt(4)
        auth_run = auth_p.add_run("Nama Penulis 1, Nama Penulis 2, Nama Penulis 3")
        auth_run.font.name = 'Times New Roman'
        auth_run.font.size = Pt(10)
        auth_run.font.bold = True
        
        # Affiliation (10pt, Centered)
        aff_p = doc.add_paragraph()
        aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        aff_p.paragraph_format.space_after = Pt(18)
        aff_run = aff_p.add_run("Jurusan Teknologi Informasi, Politeknik Negeri Malang, Indonesia\nemail1@polinema.ac.id, email2@polinema.ac.id")
        aff_run.font.name = 'Times New Roman'
        aff_run.font.size = Pt(10)
        
        # Abstract Header (10pt, Bold, Centered)
        abs_h_p = doc.add_paragraph()
        abs_h_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abs_h_p.paragraph_format.space_before = Pt(12)
        abs_h_p.paragraph_format.space_after = Pt(6)
        abs_h_run = abs_h_p.add_run("Abstrak")
        abs_h_run.font.name = 'Times New Roman'
        abs_h_run.font.size = Pt(10)
        abs_h_run.font.bold = True
        
        # Abstract Body (10pt, Justified, Single Column)
        abs_p = doc.add_paragraph(
            "Abstrak memuat permasalahan yang dikaji, metode yang digunakan, tesa-tesa (jika ada) "
            "yang dikemukakan, ulasan singkat serta penjelasan hasil penelitian dan kesimpulan yang diperoleh. "
            "Abstrak terdiri dari 200-250 kata, ditulis dalam format satu kolom dengan perataan justified."
        )
        abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        abs_p.paragraph_format.left_indent = Inches(0.5)
        abs_p.paragraph_format.right_indent = Inches(0.5)
        abs_p.paragraph_format.space_after = Pt(12)
        
        # Keywords
        key_p = doc.add_paragraph()
        key_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        key_p.paragraph_format.left_indent = Inches(0.5)
        key_p.paragraph_format.right_indent = Inches(0.5)
        key_p.paragraph_format.space_after = Pt(24)
        lbl_run = key_p.add_run("Kata kunci: ")
        lbl_run.font.bold = True
        val_run = key_p.add_run("format jip, template jurnal, pemrograman python, dokumen office")
        val_run.font.italic = True
        
        # --- Section 2: Main Body (Two Columns) ---
        # Add a section break to transition into two columns
        body_section = doc.add_section(WD_SECTION.NEW_PAGE)
        body_section.top_margin = Inches(1.18)
        body_section.left_margin = Inches(1.18)
        body_section.bottom_margin = Inches(0.78)
        body_section.right_margin = Inches(0.78)
        
        # Set two columns
        sectPr = body_section._sectPr
        cols = sectPr.xpath('./w:cols')
        if cols:
            cols[0].set(qn('w:num'), '2')
            cols[0].set(qn('w:space'), '720')  # Space between columns (approx 0.5 in)
            
        # Pendahuluan
        p_h1 = doc.add_paragraph()
        p_h1_run = p_h1.add_run("1. Pendahuluan")
        p_h1_run.font.name = 'Times New Roman'
        p_h1_run.font.size = Pt(10)
        p_h1_run.font.bold = True
        p_h1.paragraph_format.space_before = Pt(12)
        p_h1.paragraph_format.space_after = Pt(6)
        
        p_body = doc.add_paragraph(
            "Pendahuluan berisi hal-hal tentang deskripsi topik penelitian, latar belakang, "
            "masalah penelitian, tujuan, dan lingkup permasalahan. Penjelasan penelitian terkait "
            "atau terdahulu disertakan untuk menunjukkan kontribusi ilmiah naskah ini."
        )
        p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_body.paragraph_format.line_spacing = 1.15
        p_body.paragraph_format.space_after = Pt(12)
        
        # Metode
        m_h1 = doc.add_paragraph()
        m_h1_run = m_h1.add_run("2. Metode")
        m_h1_run.font.name = 'Times New Roman'
        m_h1_run.font.size = Pt(10)
        m_h1_run.font.bold = True
        m_h1.paragraph_format.space_before = Pt(12)
        m_h1.paragraph_format.space_after = Pt(6)
        
        m_body = doc.add_paragraph(
            "Pada bagian metode sebaiknya dicantumkan langkah-langkah dalam melakukan penelitian secara jelas. "
            "Memuat juga subjek penelitian, sumber data, dan alur pemrosesan data."
        )
        m_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        m_body.paragraph_format.line_spacing = 1.15
        m_body.paragraph_format.space_after = Pt(12)
        
        # Table (Label ABOVE table, 8pt)
        t_lbl = doc.add_paragraph()
        t_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t_lbl_run = t_lbl.add_run("Tabel 1. Hasil Pengujian Hyperparameter")
        t_lbl_run.font.name = 'Times New Roman'
        t_lbl_run.font.size = Pt(8)
        t_lbl_run.font.bold = True
        t_lbl.paragraph_format.space_after = Pt(4)
        
        table_data = [
            ["Model", "Duration", "Status"],
            ["LR Tuning", "4.84s", "Complete"],
            ["RF Tuning", "12.11s", "Complete"]
        ]
        
        table = doc.add_table(rows=3, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Normal Table'
        
        tblPr = table._tbl.tblPr
        tblBorders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(tblBorders)
        
        for r_idx, row in enumerate(table.rows):
            for c_idx, val in enumerate(table_data[r_idx]):
                cell = row.cells[c_idx]
                cell.text = val
                cell.width = Inches(1.5)
                set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
                
                cell_p = cell.paragraphs[0]
                cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_p.paragraph_format.space_before = Pt(0)
                cell_p.paragraph_format.space_after = Pt(0)
                for run in cell_p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(8)
                    if r_idx == 0:
                        run.font.bold = True
                        
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
        # Kesimpulan
        k_h1 = doc.add_paragraph()
        k_h1_run = k_h1.add_run("3. Kesimpulan")
        k_h1_run.font.name = 'Times New Roman'
        k_h1_run.font.size = Pt(10)
        k_h1_run.font.bold = True
        k_h1.paragraph_format.space_before = Pt(12)
        k_h1.paragraph_format.space_after = Pt(6)
        
        k_body = doc.add_paragraph(
            "Kesimpulan harus dituliskan dalam bentuk paragraf yang padat dan jelas tanpa format numbering, "
            "serta menyampaikan usulan kelanjutan riset masa depan (future works)."
        )
        k_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        k_body.paragraph_format.line_spacing = 1.15
        k_body.paragraph_format.space_after = Pt(12)
        
        # Save Document
        doc.save(filename)
        print(f"JIP formatted template successfully created: {filename}")
    except Exception as e:
        print(f"Error creating JIP template: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Create a Word document template for JIP.")
    parser.add_argument("-o", "--output", default="template_jip_filled.docx", help="Filename of the JIP document to generate.")
    args = parser.parse_args()
    create_jip_template(args.output)
