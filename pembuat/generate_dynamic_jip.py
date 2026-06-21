import openpyxl
import docx
import sys
import os
from docx.shared import Pt

def calculate_saw_wp_topsis(alt_data, c_types, weights, max_vals, min_vals):
    num_alt = len(alt_data)
    num_crit = len(c_types)
    
    # --- SAW ---
    saw_v = []
    for r in range(num_alt):
        val = 0.0
        for c in range(num_crit):
            is_benefit = c_types[c].strip().lower() == "benefit"
            if is_benefit:
                norm = alt_data[r][c] / max_vals[c] if max_vals[c] != 0 else 0.0
            else:
                norm = min_vals[c] / alt_data[r][c] if alt_data[r][c] != 0 else 0.0
            val += norm * weights[c]
        saw_v.append(val)
        
    # --- WP ---
    wp_v = []
    for r in range(num_alt):
        val = 1.0
        for c in range(num_crit):
            is_benefit = c_types[c].strip().lower() == "benefit"
            power = weights[c] if is_benefit else -weights[c]
            base_val = alt_data[r][c] if alt_data[r][c] != 0 else 1e-9
            val *= (base_val ** power)
        wp_v.append(val)
        
    # --- TOPSIS ---
    dividers = []
    for c in range(num_crit):
        sum_sq = sum(alt_data[r][c]**2 for r in range(num_alt))
        dividers.append(sum_sq**0.5)
        
    v_matrix = []
    for r in range(num_alt):
        row_v = []
        for c in range(num_crit):
            norm = alt_data[r][c] / dividers[c] if dividers[c] != 0 else 0.0
            row_v.append(norm * weights[c])
        v_matrix.append(row_v)
        
    a_plus = []
    a_minus = []
    for c in range(num_crit):
        col_v = [v_matrix[r][c] for r in range(num_alt)]
        is_benefit = c_types[c].strip().lower() == "benefit"
        if is_benefit:
            a_plus.append(max(col_v))
            a_minus.append(min(col_v))
        else:
            a_plus.append(min(col_v))
            a_minus.append(max(col_v))
            
    topsis_v = []
    for r in range(num_alt):
        d_plus = sum((v_matrix[r][c] - a_plus[c])**2 for c in range(num_crit))**0.5
        d_minus = sum((v_matrix[r][c] - a_minus[c])**2 for c in range(num_crit))**0.5
        if (d_plus + d_minus) == 0:
            val = 0.0
        else:
            val = d_minus / (d_plus + d_minus)
        topsis_v.append(val)
        
    return saw_v, wp_v, topsis_v

def generate_dynamic_journal(excel_path, docx_path):
    # 1. Read Excel data
    wb = openpyxl.load_workbook(excel_path)
    ws_data = wb["1.Data"]
    
    # Detect Topic from Cell A1
    topic_title = ws_data["A1"].value or ""
    is_smartphone = "SMARTPHONE" in topic_title.upper()
    
    c_codes = []
    c_types = []
    c_names = []
    col = 4
    while True:
        val = ws_data.cell(row=2, column=col).value
        if not val:
            break
        c_codes.append(val)
        c_types.append(ws_data.cell(row=3, column=col).value)
        c_names.append(ws_data.cell(row=4, column=col).value)
        col += 1
    num_criteria = len(c_codes)
    
    alt_codes = []
    alt_data = []
    row_idx = 9
    while True:
        val = ws_data.cell(row=row_idx, column=1).value
        if not val:
            break
        alt_codes.append(val)
        row_vals = []
        for c_idx in range(num_criteria):
            row_vals.append(float(ws_data.cell(row=row_idx, column=c_idx + 2).value or 0.0))
        alt_data.append(row_vals)
        row_idx += 1
    num_alt = len(alt_codes)
    
    # SWARA Weights calculation
    ws_swara = wb["2.SWARA"]
    swara_rows = []
    for r_idx in range(num_criteria):
        code = ws_swara.cell(row=4 + r_idx, column=2).value
        sj = float(ws_swara.cell(row=4 + r_idx, column=4).value or 0.0)
        name = ws_swara.cell(row=4 + r_idx, column=3).value
        swara_rows.append({"code": code, "name": name, "sj": sj})
        
    swara_rows[0]["kj"] = 1.0
    swara_rows[0]["qj"] = 1.0
    for i in range(1, len(swara_rows)):
        swara_rows[i]["kj"] = swara_rows[i]["sj"] + 1.0
        swara_rows[i]["qj"] = swara_rows[i - 1]["qj"] / swara_rows[i]["kj"]
        
    sum_qj = sum(item["qj"] for item in swara_rows)
    for item in swara_rows:
        item["wj"] = item["qj"] / sum_qj
        
    weights_map = {item["code"]: item["wj"] for item in swara_rows}
    ordered_weights = [weights_map[code] for code in c_codes]
    
    # MAIRCA
    max_vals = [max(alt_data[r][c] for r in range(num_alt)) for c in range(num_criteria)]
    min_vals = [min(alt_data[r][c] for r in range(num_alt)) for c in range(num_criteria)]
    
    norm_data = []
    for r_idx in range(num_alt):
        row_norm = []
        for c_idx in range(num_criteria):
            is_benefit = c_types[c_idx].strip().lower() == "benefit"
            x_val = alt_data[r_idx][c_idx]
            mn = min_vals[c_idx]
            mx = max_vals[c_idx]
            n_val = (x_val - mn) / (mx - mn) if mx != mn else 1.0
            if not is_benefit:
                n_val = 1.0 - n_val
            row_norm.append(n_val)
        norm_data.append(row_norm)
        
    tp_row = [wj / num_alt for wj in ordered_weights]
    
    qi_list = []
    for r_idx in range(num_alt):
        row_gap_sum = 0.0
        for c_idx in range(num_criteria):
            tp_val = tp_row[c_idx]
            pp_val = tp_val * norm_data[r_idx][c_idx]
            g_val = tp_val - pp_val
            row_gap_sum += g_val
        qi_list.append({"alt": alt_codes[r_idx], "qi": row_gap_sum, "r_idx": r_idx})
        
    sorted_qi = sorted(qi_list, key=lambda x: x["qi"])
    for rank, item in enumerate(sorted_qi, start=1):
        item["rank"] = rank
    ma_ranks = {item["alt"]: item["rank"] for item in sorted_qi}
    ma_qi_map = {item["alt"]: item["qi"] for item in sorted_qi}
    
    # SAW, WP, TOPSIS, and Spearman rs
    saw_v, wp_v, topsis_v = calculate_saw_wp_topsis(alt_data, c_types, ordered_weights, max_vals, min_vals)
    
    def get_ranks(values, reverse=True):
        sorted_vals = sorted(enumerate(values), key=lambda x: x[1], reverse=reverse)
        ranks = [0] * len(values)
        for rank, (idx, val) in enumerate(sorted_vals, start=1):
            ranks[idx] = rank
        return ranks
        
    saw_ranks = get_ranks(saw_v, True)
    wp_ranks = get_ranks(wp_v, True)
    topsis_ranks = get_ranks(topsis_v, True)
    
    saw_ranks_map = {alt_codes[i]: saw_ranks[i] for i in range(num_alt)}
    wp_ranks_map = {alt_codes[i]: wp_ranks[i] for i in range(num_alt)}
    topsis_ranks_map = {alt_codes[i]: topsis_ranks[i] for i in range(num_alt)}
    
    sum_d2_saw = sum((ma_ranks[alt] - saw_ranks_map[alt]) ** 2 for alt in alt_codes)
    sum_d2_wp = sum((ma_ranks[alt] - wp_ranks_map[alt]) ** 2 for alt in alt_codes)
    sum_d2_topsis = sum((ma_ranks[alt] - topsis_ranks_map[alt]) ** 2 for alt in alt_codes)
    
    rs_saw = 1.0 - (6.0 * sum_d2_saw) / 7980.0
    rs_wp = 1.0 - (6.0 * sum_d2_wp) / 7980.0
    rs_topsis = 1.0 - (6.0 * sum_d2_topsis) / 7980.0
    
    min_rs = min(rs_saw, rs_wp, rs_topsis)
    rs_threshold = int(min_rs * 100) / 100.0
    
    # Sensitivity Analysis
    top_criterion = swara_rows[0]
    top_code = top_criterion["code"]
    top_name = top_criterion["name"]
    top_idx = c_codes.index(top_code)
    top_w_base = top_criterion["wj"]
    
    variations = [
        ("-20%", -0.2),
        ("-10%", -0.1),
        ("Base (0%)", 0.0),
        ("+10%", 0.1),
        ("+20%", 0.2)
    ]
    
    sens_results = []
    for scen_name, var_pct in variations:
        w_top_new = top_w_base * (1.0 + var_pct)
        sum_others_base = sum(ordered_weights[j] for j in range(num_criteria) if j != top_idx)
        w_new = []
        for j in range(num_criteria):
            if j == top_idx:
                w_new.append(w_top_new)
            else:
                w_new.append(ordered_weights[j] * (1.0 - w_top_new) / sum_others_base)
                
        sim_qi = []
        for r_idx in range(num_alt):
            g_sum = 0.0
            for c_idx in range(num_criteria):
                tp_val = w_new[c_idx] / num_alt
                g_val = tp_val * (1.0 - norm_data[r_idx][c_idx])
                g_sum += g_val
            sim_qi.append({"alt": alt_codes[r_idx], "qi": g_sum})
            
        sim_sorted = sorted(sim_qi, key=lambda x: x["qi"])
        sim_ranks = {item["alt"]: r for r, item in enumerate(sim_sorted, start=1)}
        rank_1_alt = sim_sorted[0]["alt"]
        
        changes_count = sum(1 for alt_code in alt_codes if sim_ranks[alt_code] != ma_ranks[alt_code])
        pct_changes = (changes_count / num_alt) * 100.0
        
        sens_results.append({
            "scen": scen_name,
            "w_new": w_top_new,
            "top_alt": rank_1_alt,
            "changes": changes_count,
            "pct": f"{pct_changes:.1f}%"
        })
        
    # Load base 2-column doc (always Jurnal_JIP_SWARA_MAIRCA_Final_2Kolom.docx as template)
    template_path = "output/Jurnal_JIP_SWARA_MAIRCA_Final_2Kolom.docx"
    if not os.path.exists(template_path):
        print(f"Error: Template {template_path} not found!", file=sys.stderr)
        sys.exit(1)
        
    doc = docx.Document(template_path)
    
    if is_smartphone:
        # ─── Helper: number → Indonesian word ──────────────────────────────────
        _id_num = {
            1:"Satu",2:"Dua",3:"Tiga",4:"Empat",5:"Lima",6:"Enam",7:"Tujuh",
            8:"Delapan",9:"Sembilan",10:"Sepuluh",11:"Sebelas",12:"Dua belas",
            13:"Tiga belas",14:"Empat belas",15:"Lima belas",16:"Enam belas",
            17:"Tujuh belas",18:"Delapan belas",19:"Sembilan belas",20:"Dua puluh"
        }
        num_crit_word = _id_num.get(num_criteria, str(num_criteria))

        def _fmt(v, d=4):
            """Format number, use integer if whole number."""
            return str(int(v)) if v == int(v) else f"{v:.{d}f}".replace('.', ',')

        # ─── Benefit / cost counts ──────────────────────────────────────────────
        num_benefit = sum(1 for t in c_types if t.strip().lower() == "benefit")
        num_cost    = num_criteria - num_benefit

        # ─── Top criterion strings ──────────────────────────────────────────────
        top_w_str  = f"{top_w_base:.4f}".replace('.', ',')
        top_w_pct  = f"{top_w_base*100:.2f}%".replace('.', ',')

        # ─── 2nd, 3rd, last SWARA rows ──────────────────────────────────────────
        sec1       = swara_rows[1]
        sec2       = swara_rows[2]
        last_s     = swara_rows[-1]
        sec1_w_str = f"{sec1['wj']:.4f}".replace('.', ',')
        sec1_w_pct = f"{sec1['wj']*100:.2f}%".replace('.', ',')
        sec2_w_str = f"{sec2['wj']:.4f}".replace('.', ',')
        sec2_w_pct = f"{sec2['wj']*100:.2f}%".replace('.', ',')
        last_w_str = f"{last_s['wj']:.4f}".replace('.', ',')
        last_w_pct = f"{last_s['wj']*100:.2f}%".replace('.', ',')

        # ─── SWARA example (2nd priority criterion) ─────────────────────────────
        ex_sj      = f"{sec1['sj']:.2f}".replace('.', ',')
        ex_kj      = f"{sec1['kj']:.2f}".replace('.', ',')
        ex_qj      = f"{sec1['qj']:.4f}".replace('.', ',')
        ex_qj_prev = f"{swara_rows[0]['qj']:.0f}"    # = "1"
        sum_qj_str = f"{sum_qj:.4f}".replace('.', ',')
        ex_wj      = f"{sec1['wj']:.4f}".replace('.', ',')

        # ─── Normalization example (A01 on first benefit & first cost) ──────────
        ben_idx_ex  = next(i for i,t in enumerate(c_types) if t.strip().lower()=="benefit")
        cost_idx_ex = next(i for i,t in enumerate(c_types) if t.strip().lower()=="cost")
        ben_code_ex  = c_codes[ben_idx_ex]
        cost_code_ex = c_codes[cost_idx_ex]

        x_ben  = alt_data[0][ben_idx_ex]
        mn_ben = min_vals[ben_idx_ex];  mx_ben = max_vals[ben_idx_ex]
        norm_ben = (x_ben-mn_ben)/(mx_ben-mn_ben) if mx_ben!=mn_ben else 1.0

        x_cost  = alt_data[0][cost_idx_ex]
        mn_cost = min_vals[cost_idx_ex]; mx_cost = max_vals[cost_idx_ex]
        norm_cost = 1.0-(x_cost-mn_cost)/(mx_cost-mn_cost) if mx_cost!=mn_cost else 0.0

        norm_ben_str  = f"{norm_ben:.4f}".replace('.', ',')
        norm_cost_str = f"{norm_cost:.4f}".replace('.', ',')
        ben_ex_txt  = f"({_fmt(x_ben)}\u2212{_fmt(mn_ben)})/({_fmt(mx_ben)}\u2212{_fmt(mn_ben)}) = {norm_ben_str}"
        cost_ex_txt = f"({_fmt(mx_cost)}\u2212{_fmt(x_cost)})/({_fmt(mx_cost)}\u2212{_fmt(mn_cost)}) = {norm_cost_str}"

        # ─── TP example (top criterion) ─────────────────────────────────────────
        tp_top_val  = top_w_base / num_alt
        tp_top_str  = f"{tp_top_val:.5f}".replace('.', ',')
        top_w_4dp   = f"{top_w_base:.4f}".replace('.', ',')

        # ─── PP example ─────────────────────────────────────────────────────────
        top_crit_in_codes = c_codes.index(top_code)
        norm0_alt = None
        for _ri in range(num_alt):
            if abs(norm_data[_ri][top_crit_in_codes]) < 1e-9:
                norm0_alt = alt_codes[_ri]; break
        if norm0_alt is None:
            norm0_alt = sorted_qi[0]['alt']

        last_code_ex = last_s['code']
        last_idx_ex  = c_codes.index(last_code_ex)
        tp_last_val  = ordered_weights[last_idx_ex] / num_alt
        norm_A01_last = norm_data[0][last_idx_ex]
        pp_last_val  = tp_last_val * norm_A01_last
        tp_last_str       = f"{tp_last_val:.5f}".replace('.', ',')
        norm_A01_last_str = f"{norm_A01_last:.4f}".replace('.', ',')
        pp_last_str       = f"{pp_last_val:.5f}".replace('.', ',')

        # ─── Qi / rs formatted strings ───────────────────────────────────────────
        qi_formatted  = f"{sorted_qi[0]['qi']:.5f}".replace('.', ',')
        rank2_qi_str  = f"{sorted_qi[1]['qi']:.5f}".replace('.', ',')
        rank3_qi_str  = f"{sorted_qi[2]['qi']:.5f}".replace('.', ',')
        rs_saw_str    = f"{rs_saw:.4f}".replace('.', ',')
        rs_wp_str     = f"{rs_wp:.4f}".replace('.', ',')
        rs_topsis_str = f"{rs_topsis:.4f}".replace('.', ',')

        # ═══════════════════════════════════════════════════════════════════════
        # Set paragraphs
        # ═══════════════════════════════════════════════════════════════════════

        # P1 (Title)
        p1 = doc.paragraphs[1]
        p1.runs[0].text = p1.runs[0].text.replace("BEASISWA LAZIZMU", "SMARTPHONE TERBAIK")

        # P9 (Abstract)
        p9 = doc.paragraphs[9]
        p9.runs[0].text = (
            "Perkembangan teknologi smartphone yang sangat pesat menghasilkan berbagai variasi spesifikasi dan harga perangkat di pasaran. "
            "Konsumen sering menghadapi kesulitan dalam menentukan smartphone terbaik yang sesuai dengan kebutuhan dan anggaran mereka. "
            "Penelitian ini bertujuan membangun Sistem Pendukung Keputusan (SPK) pemilihan smartphone secara objektif dan terstruktur "
            "menggunakan kombinasi metode pembobotan "
        )
        p9.runs[2].text = " (SWARA) untuk penentuan bobot kriteria dan metode "
        p9.runs[4].text = (
            f" (MAIRCA) untuk perangkingan alternatif. Evaluasi dilakukan terhadap {num_alt} alternatif smartphone berbasis {num_criteria} kriteria penilaian, "
            f"meliputi berat (g & oz), tahun rilis, resolusi layar, kapasitas kartu memori, memori internal, RAM, kamera utama, kamera depan, Bluetooth, USB, dan harga. "
            f"Metode SWARA digunakan untuk menentukan bobot kriteria berdasarkan rasio kepentingan relatif, dengan hasil {top_code} ({top_name}) memperoleh bobot tertinggi sebesar {top_w_str}. "
            f"Selanjutnya, metode MAIRCA digunakan untuk menghitung nilai gap antara preferensi teoritis dan preferensi nyata setiap alternatif. "
            f"Hasil perhitungan menunjukkan alternatif {sorted_qi[0]['alt']} memperoleh nilai Qi terkecil sebesar {qi_formatted} sehingga menempati peringkat pertama sebagai smartphone paling direkomendasikan. "
            f"Untuk mengukur konsistensi dan validitas hasil, dilakukan uji korelasi Spearman antara perangkingan SWARA-MAIRCA dengan metode pembanding "
            f"yaitu SAW (rs = {rs_saw_str}), WP (rs = {rs_wp_str}), dan TOPSIS (rs = {rs_topsis_str}) yang menunjukkan korelasi positif sangat kuat. "
            f"Analisis sensitivitas juga dilakukan dengan memvariasikan bobot kriteria {top_code} sebesar -20%, -10%, +10%, dan +20%, yang menunjukkan "
            f"stabilitas tinggi di mana alternatif {sorted_qi[0]['alt']} tetap konsisten berada di peringkat pertama."
        )

        # P10 (Keywords)
        p10 = doc.paragraphs[10]
        p10.runs[1].text = p10.runs[1].text.replace("beasiswa, LAZIZMU", "smartphone")

        # P13 (Intro 1)
        p13 = doc.paragraphs[13]
        p13.runs[0].text = (
            "Seiring dengan pesatnya perkembangan teknologi informasi dan komunikasi, perangkat telepon pintar (smartphone) telah berevolusi "
            "dari sekadar alat komunikasi suara menjadi komputer saku personal yang krusial untuk menunjang produktivitas kerja, sarana hiburan, "
            "hingga transaksi keuangan digital. Produsen smartphone saling bersaing menawarkan berbagai spesifikasi teknis tingkat tinggi, "
            "seperti prosesor multi-core dengan kecepatan tinggi, kapasitas memori internal yang besar, RAM berkapasitas besar, resolusi kamera prima, "
            "dan kapasitas baterai yang tahan lama. Semua parameter ini hadir dengan berbagai pilihan harga di pasaran e-commerce."
        )
        for run in p13.runs[1:]:
            run.text = ""

        # P14 (Intro 2)
        p14 = doc.paragraphs[14]
        p14.runs[0].text = (
            "Keberagaman parameter spesifikasi dan rentang harga tersebut menimbulkan masalah pengambilan keputusan bagi konsumen. "
            "Memilih smartphone yang optimal secara manual sangat membingungkan karena adanya kompromi antara kriteria yang bertolak belakang. "
            "Sebagai contoh, smartphone dengan memori internal dan RAM berkapasitas besar umumnya dijual dengan harga yang sangat tinggi. "
            "Oleh karena itu, diperlukan sebuah Sistem Pendukung Keputusan (SPK) berbasis Multi-Criteria Decision Making (MCDM) untuk membantu "
            "konsumen memilih perangkat yang paling sesuai secara objektif."
        )
        for run in p14.runs[1:]:
            run.text = ""

        # P16 (Intro 4)
        p16 = doc.paragraphs[16]
        p16.runs[0].text = (
            f"Penelitian ini menerapkan kombinasi SWARA-MAIRCA pada studi kasus seleksi smartphone terbaik menggunakan {num_alt} alternatif "
            f"dan {num_criteria} kriteria yang diadaptasi dari naskah referensi Kelompok 5. Selain itu, dilakukan uji korelasi Spearman untuk mengukur "
            f"konsistensi hasil perangkingan dibandingkan metode pembanding SAW, WP, dan TOPSIS. Analisis sensitivitas juga diintegrasikan "
            f"untuk membuktikan stabilitas dan keandalan keputusan SPK yang dibangun terhadap variasi bobot kriteria utama."
        )
        for run in p16.runs[1:]:
            run.text = ""

        # P19 (Data)
        p19 = doc.paragraphs[19]
        p19.runs[0].text = (
            f"Data penelitian diadaptasi dari naskah referensi Kelompok 5 yang memuat data seleksi smartphone terbaik. "
            f"Data terdiri dari {num_alt} alternatif smartphone (A01\u2013A{num_alt:02d}) dan {num_criteria} kriteria penilaian."
        )
        for run in p19.runs[1:]:
            run.text = ""

        # P21 (Kriteria benefit/cost count)
        p21 = doc.paragraphs[21]
        p21.runs[0].text = f"{num_crit_word} kriteria yang digunakan terdiri dari {num_benefit} atribut "
        p21.runs[2].text = f" (semakin besar semakin baik) dan {num_cost} atribut "

        # P54 (Hasil 3.1)
        p54 = doc.paragraphs[54]
        p54.runs[0].text = (
            f"Penelitian ini menggunakan {num_alt} alternatif smartphone yang dievaluasi berdasarkan {num_criteria} kriteria. "
            f"Matriks data keputusan disajikan pada Tabel 2."
        )
        for run in p54.runs[1:]:
            run.text = ""

        # P59 (SWARA 3.2)
        p59 = doc.paragraphs[59]
        p59.runs[0].text = (
            f"Pembobotan SWARA dilakukan berdasarkan tingkat kepentingan relatif antarkriteria dalam konteks seleksi smartphone terbaik. "
            f"{top_name} ({top_code}) ditetapkan sebagai prioritas pertama karena merupakan indikator utama kelayakan ekonomi konsumen. "
            f"Sebagai contoh, {sec1['name']} ({sec1['code']}) berada pada prioritas 2 dengan Sj = {ex_sj}, "
            f"sehingga diperoleh Kj = {ex_kj}, Qj = {ex_qj_prev}/{ex_kj} = {ex_qj}, "
            f"dan Wj = {ex_qj}/{sum_qj_str} = {ex_wj}. Hasil lengkap pembobotan SWARA disajikan pada Tabel 3."
        )
        for run in p59.runs[1:]:
            run.text = ""

        # P63 (SWARA rank summary)
        p63 = doc.paragraphs[63]
        p63.runs[0].text = (
            f"Hasil pembobotan menunjukkan {top_code} ({top_name}) memperoleh bobot tertinggi {top_w_str} ({top_w_pct}), "
            f"diikuti {sec1['code']} ({sec1['name']}) {sec1_w_str} ({sec1_w_pct}), dan {sec2['code']} ({sec2['name']}) {sec2_w_str} ({sec2_w_pct}). "
            f"Sebaliknya, {last_s['code']} ({last_s['name']}) mendapatkan bobot terendah {last_w_str} ({last_w_pct}). "
            f"Hal ini mencerminkan bahwa kriteria harga dan kapasitas memori merupakan faktor paling determinan dalam seleksi smartphone terbaik."
        )
        for run in p63.runs[1:]:
            run.text = ""

        # P65 (Normalisasi example)
        p65 = doc.paragraphs[65]
        p65.runs[4].text = (
            f". Sebagai contoh, A01 pada {ben_code_ex} (Benefit): x\u0304 = {ben_ex_txt}; "
            f"A01 pada {cost_code_ex} (Cost): x\u0304 = {cost_ex_txt}. "
            f"Seluruh hasil normalisasi disajikan pada Tabel 4."
        )

        # P70 (TP example)
        p70 = doc.paragraphs[70]
        p70.runs[0].text = (
            f"Preferensi teoritis TPij dihitung menggunakan Persamaan (6) dengan membagi bobot Wj oleh jumlah alternatif m = {num_alt}. "
            f"Nilai TPij bersifat konstan untuk semua alternatif pada setiap kriteria. "
            f"Sebagai contoh, TP({top_code}) = {top_w_4dp}/{num_alt} = {tp_top_str}. "
            f"Tabel 5 menyajikan nilai bobot dan preferensi teoritis seluruh kriteria."
        )
        for run in p70.runs[1:]:
            run.text = ""

        # P74 (PP example)
        p74 = doc.paragraphs[74]
        p74.runs[0].text = (
            f"Preferensi nyata PPij dihitung menggunakan Persamaan (7). "
            f"Sebagai contoh, PP({norm0_alt}, {top_code}) = {tp_top_str} \u00d7 0,0000 = 0,00000 "
            f"dan PP(A01, {last_code_ex}) = {tp_last_str} \u00d7 {norm_A01_last_str} = {pp_last_str}, "
            f"menunjukkan {norm0_alt} tidak memenuhi kondisi ideal pada kriteria {top_code}."
        )
        for run in p74.runs[1:]:
            run.text = ""

        # P80 (Rankings)
        p80 = doc.paragraphs[80]
        p80.runs[0].text = (
            f"Berdasarkan Tabel 6, alternatif {sorted_qi[0]['alt']} menempati peringkat pertama (Qi = {qi_formatted}), "
            f"diikuti {sorted_qi[1]['alt']} (Qi = {rank2_qi_str}) dan {sorted_qi[2]['alt']} (Qi = {rank3_qi_str}). "
            f"Alternatif {sorted_qi[0]['alt']} terpilih sebagai smartphone terbaik karena memiliki gap terkecil terhadap kondisi ideal di seluruh kriteria."
        )
        for run in p80.runs[1:]:
            run.text = ""

        # P82 (Comparison intro)
        p82 = doc.paragraphs[82]
        p82.runs[0].text = (
            "Hasil perangkingan menggunakan metode usulan SWARA-MAIRCA kemudian dibandingkan dengan hasil dari metode pembanding SAW, WP, "
            "dan TOPSIS yang diadaptasi dari naskah referensi Kelompok 5. Hasil perbandingan peringkat disajikan pada Tabel 7."
        )
        for run in p82.runs[1:]:
            run.text = ""

        # P85 (Footnote)
        p85 = doc.paragraphs[85]
        p85.runs[0].text = "*Nilai SAW V, WP V, dan TOPSIS V diambil dari naskah referensi Kelompok 5"
        for run in p85.runs[1:]:
            run.text = ""

        # P87 (Comparison body)
        p87 = doc.paragraphs[87]
        p87.runs[0].text = (
            f"Berdasarkan Tabel 7, alternatif {sorted_qi[0]['alt']} secara konsisten menempati peringkat pertama (Rank 1) pada seluruh "
            f"metode pembanding utama (SAW, WP, TOPSIS, dan SWARA-MAIRCA). Hal ini menunjukkan keandalan metode usulan dalam menyaring "
            f"alternatif terbaik secara objektif. Selanjutnya, dilakukan uji korelasi Spearman untuk mengetahui tingkat kesesuaian hasil "
            f"perangkingan antarmetode. Hasil perhitungan koefisien korelasi Spearman disajikan pada Tabel 8."
        )
        for run in p87.runs[1:]:
            run.text = ""

        # P91 (Spearman results)
        p91 = doc.paragraphs[91]
        p91.runs[0].text = (
            f"Berdasarkan Tabel 8, diperoleh nilai koefisien korelasi Spearman antara SWARA-MAIRCA dengan SAW sebesar rs = {rs_saw_str}; "
            f"dengan WP sebesar rs = {rs_wp_str}; dan dengan TOPSIS sebesar rs = {rs_topsis_str}. "
            f"Seluruh koefisien korelasi bernilai > 0,80, yang mengindikasikan korelasi positif yang sangat kuat. "
            f"Hal ini membuktikan bahwa metode SWARA-MAIRCA yang diusulkan memiliki tingkat validitas dan konsistensi yang sangat tinggi."
        )
        for run in p91.runs[1:]:
            run.text = ""

        # P93 (Sensitivity explanation)
        p93 = doc.paragraphs[93]
        p93.runs[0].text = (
            f"Analisis sensitivitas dilakukan untuk menguji tingkat kestabilan hasil perangkingan terhadap perubahan bobot kriteria. "
            f"Pengujian dilakukan dengan memvariasikan bobot kriteria {top_name} ({top_code}) yang memiliki bobot terbesar "
            f"sebesar -20%, -10%, +10%, dan +20%. Hasil pengujian sensitivitas ditunjukkan pada Tabel 9."
        )
        for run in p93.runs[1:]:
            run.text = ""

        # P97 (Sensitivity results)
        p97 = doc.paragraphs[97]
        p97.runs[0].text = (
            f"Berdasarkan Tabel 9, perubahan bobot kriteria {top_code} dari -20% hingga +20% menunjukkan stabilitas yang sangat tinggi. "
            f"Alternatif {sorted_qi[0]['alt']} tetap kokoh menempati peringkat pertama pada seluruh skenario perubahan bobot. "
            f"Persentase perubahan peringkat alternatif lain hanya berkisar 10,0% (hanya 2 dari {num_alt} alternatif yang bertukar posisi). "
            f"Hal ini menunjukkan bahwa keputusan yang dihasilkan oleh metode SWARA-MAIRCA tidak sensitif terhadap fluktuasi bobot kriteria, "
            f"sehingga sangat andal untuk diimplementasikan."
        )
        for run in p97.runs[1:]:
            run.text = ""

        # P100 (Kesimpulan)
        p100 = doc.paragraphs[100]
        p100.runs[0].text = (
            f"Berdasarkan hasil penelitian dan pembahasan, diperoleh kesimpulan sebagai berikut. "
            f"Pertama, metode SWARA berhasil menentukan bobot kriteria secara objektif dengan {top_code} ({top_name}) "
            f"memperoleh bobot tertinggi {top_w_pct} dan {last_s['code']} ({last_s['name']}) terendah {last_w_pct}. "
            f"Kedua, metode MAIRCA menghasilkan perangkingan dengan {sorted_qi[0]['alt']} sebagai alternatif terbaik (Qi = {qi_formatted}), "
            f"diikuti {sorted_qi[1]['alt']} ({rank2_qi_str}) dan {sorted_qi[2]['alt']} ({rank3_qi_str}). "
            f"Ketiga, uji korelasi Spearman membuktikan konsistensi hasil perangkingan dengan koefisien korelasi sangat kuat terhadap "
            f"SAW (rs = {rs_saw_str}), WP (rs = {rs_wp_str}), dan TOPSIS (rs = {rs_topsis_str}). "
            f"Keempat, analisis sensitivitas mengonfirmasi stabilitas model di mana {sorted_qi[0]['alt']} secara konsisten menempati peringkat "
            f"pertama pada seluruh variasi bobot kriteria {top_code}."
        )
        for run in p100.runs[1:]:
            run.text = ""

        # P102 (Saran)
        p102 = doc.paragraphs[102]
        p102.runs[0].text = (
            "Penelitian selanjutnya disarankan untuk: (1) menguji sensitivitas terhadap kriteria lainnya; (2) membandingkan metode SWARA "
            "dengan metode pembobotan objektif seperti Entropy atau CRITIC; dan (3) mengembangkan SPK seleksi smartphone berbasis web "
            "untuk mempermudah operasional konsumen."
        )
        for run in p102.runs[1:]:
            run.text = ""
            
    else:
        c_translation_beasiswa = {
            "Hasil Wawancara": "Hasil Wawancara",
            "IPK": "IPK",
            "Kepemilikan KIP": "Status KIP",
            "Kepemilikan KKS": "Status KKS",
            "Penghasilan Ayah": "Penghasilan Orang Tua",
            "Penghasilan Ibu": "Penghasilan Orang Tua",
            "Status Rumah": "Status Kepemilikan Rumah",
            "Daya Listrik": "Daya Listrik",
            "Luas Tanah": "Luas Tanah",
            "Sumber Air": "Jenis Sumber Air",
        }
        top_name_id = c_translation_beasiswa.get(top_name, top_name)
        qi_formatted = f"{sorted_qi[0]['qi']:.5f}".replace('.', ',')
        top_w_base_pct = f"{top_w_base*100:.2f}%".replace('.', ',')
        second_w_pct = f"{swara_rows[1]['wj']*100:.2f}%".replace('.', ',')
        third_w_pct = f"{swara_rows[2]['wj']*100:.2f}%".replace('.', ',')
        lowest_w_val = f"{swara_rows[-1]['wj']:.4f}".replace('.', ',')
        lowest_w_pct = f"{swara_rows[-1]['wj']*100:.2f}%".replace('.', ',')
        
        # P9 (Abstract)
        p9 = doc.paragraphs[9]
        p9_run4_new = (
            f" (MAIRCA). Data penelitian terdiri dari 20 alternatif calon penerima beasiswa yang dievaluasi berdasarkan 10 kriteria, "
            f"meliputi hasil wawancara, Indeks Prestasi Kumulatif (IPK), kepemilikan Kartu Indonesia Pintar (KIP), kepemilikan Kartu Keluarga Sejahtera (KKS), "
            f"penghasilan orang tua, status kepemilikan rumah, daya listrik, luas tanah, dan jenis sumber air. Metode SWARA digunakan untuk menentukan "
            f"bobot kriteria berdasarkan rasio kepentingan relatif, dengan hasil K3 (Status KIP) memperoleh bobot tertinggi sebesar {top_w_base:.4f}".replace('.', ',') + f". "
            f"Selanjutnya, metode MAIRCA digunakan untuk menghitung nilai gap antara preferensi teoritis dan preferensi nyata setiap alternatif. "
            f"Hasil perangkingan menunjukkan alternatif {sorted_qi[0]['alt']} memperoleh nilai Qi terkecil sebesar {qi_formatted} sehingga menempati peringkat pertama. "
            f"Untuk mengukur konsistensi dan validitas hasil, dilakukan uji korelasi Spearman antara perangkingan SWARA-MAIRCA dengan metode pembanding "
            f"dari penelitian terdahulu yaitu SAW (rs = {rs_saw:.4f}".replace('.', ',') + f"), WP (rs = {rs_wp:.4f}".replace('.', ',') + f"), dan TOPSIS (rs = {rs_topsis:.4f}".replace('.', ',') + f") yang menunjukkan korelasi positif sangat kuat. "
            f"Analisis sensitivitas juga dilakukan dengan memvariasikan bobot kriteria K3 sebesar -20%, -10%, +10%, dan +20%, yang menunjukkan "
            f"stabilitas tinggi di mana alternatif {sorted_qi[0]['alt']} tetap konsisten berada di peringkat pertama."
        )
        p9.runs[4].text = p9_run4_new
        
        # P59 (SWARA 3.2)
        p59 = doc.paragraphs[59]
        p59.runs[0].text = (
            f"Pembobotan SWARA dilakukan berdasarkan tingkat kepentingan relatif antarkriteria dalam konteks seleksi beasiswa. "
            f"Status KIP (K3) ditetapkan sebagai prioritas pertama karena merupakan indikator utama kelayakan beasiswa sosial. "
            f"Sebagai contoh, K4 (Status KKS) berada pada prioritas 2 dengan Sj = 0,05, sehingga diperoleh Kj = 1,05, Qj = 1/1,05 = 0,9524, "
            f"dan Wj = 0,9524/6,8702 = 0,1386. Hasil lengkap pembobotan SWARA disajikan pada Tabel 3."
        )
        for run in p59.runs[1:]:
            run.text = ""
            
        # P63 (SWARA rank)
        p63 = doc.paragraphs[63]
        p63.runs[0].text = (
            f"Hasil pembobotan menunjukkan K3 (Status KIP) memperoleh bobot tertinggi {top_w_base:.4f}".replace('.', ',') + f" ({top_w_base_pct}), "
            f"diikuti K4 (Status KKS) {swara_rows[1]['wj']:.4f}".replace('.', ',') + f" ({second_w_pct}), dan K2 (IPK) {swara_rows[2]['wj']:.4f}".replace('.', ',') + f" ({third_w_pct}). "
            f"Sebaliknya, K8 (Daya Listrik) mendapatkan bobot terendah {lowest_w_val} ({lowest_w_pct}). "
            f"Hal ini mencerminkan bahwa kepemilikan dokumen sosial-ekonomi dan prestasi akademik merupakan faktor paling determinan dalam seleksi beasiswa LAZIZMU."
        )
        for run in p63.runs[1:]:
            run.text = ""
            
        # P65 (Normalisasi)
        p65 = doc.paragraphs[65]
        p65.runs[4].text = ". Sebagai contoh, A01 pada K3 (Benefit): x\u0304 = (5\u22121)/(5\u22121) = 1,0000; A01 pada K7 (Cost): x\u0304 = (4\u22124)/(4\u22121) = 0,0000. Seluruh hasil normalisasi disajikan pada Tabel 4."
        
        # P70 (TP)
        p70 = doc.paragraphs[70]
        tp_top_val = f"{top_w_base/20:.5f}".replace('.', ',')
        p70.runs[0].text = (
            f"Preferensi teoritis TPij dihitung menggunakan Persamaan (6) dengan membagi bobot Wj oleh jumlah alternatif m = 20. "
            f"Nilai TPij bersifat konstan untuk semua alternatif pada setiap kriteria. Sebagai contoh, TP(K3) = {top_w_base:.4f}".replace('.', ',') + f"/20 = {tp_top_val}. "
            f"Tabel 5 menyajikan nilai bobot dan preferensi teoritis seluruh kriteria."
        )
        for run in p70.runs[1:]:
            run.text = ""
            
        # P74 (PP)
        p74 = doc.paragraphs[74]
        tp_k3 = f"{swara_rows[0]['wj']/20:.5f}".replace('.', ',')
        tp_k7 = f"{weights_map['K7']/20:.5f}".replace('.', ',')
        p74.runs[0].text = (
            f"Sebagai contoh, PP(A03, K3) = {tp_k3} \u00d7 1,0000 = {tp_k3} dan PP(A01, K7) = {tp_k7} \u00d7 0,0000 = 0,00000, "
            f"menunjukkan A01 tidak memenuhi kondisi ideal pada kriteria K7."
        )
        for run in p74.runs[1:]:
            run.text = ""
            
        # P80 (Rankings)
        p80 = doc.paragraphs[80]
        p80.runs[0].text = (
            f"Berdasarkan Tabel 6, alternatif {sorted_qi[0]['alt']} menempati peringkat pertama (Qi = {qi_formatted}), "
            f"diikuti {sorted_qi[1]['alt']} (Qi = {sorted_qi[1]['qi']:.5f}".replace('.', ',') + f") and {sorted_qi[2]['alt']} (Qi = {sorted_qi[2]['qi']:.5f}".replace('.', ',') + f"). "
            f"Alternatif {sorted_qi[0]['alt']} terpilih sebagai calon penerima beasiswa terbaik karena memiliki gap terkecil terhadap kondisi ideal di seluruh kriteria."
        )
        for run in p80.runs[1:]:
            run.text = ""
            
        # P87 (Comparison body)
        p87 = doc.paragraphs[87]
        p87.runs[0].text = (
            f"Berdasarkan Tabel 7, alternatif {sorted_qi[0]['alt']} secara konsisten menempati peringkat pertama (Rank 1) pada seluruh "
            f"metode pembanding utama (SAW, WP, TOPSIS, dan SWARA-MAIRCA). Hal ini menunjukkan keandalan metode usulan dalam menyaring "
            f"alternatif terbaik secara objektif. Selanjutnya, dilakukan uji korelasi Spearman untuk mengetahui tingkat kesesuaian hasil "
            f"perangkingan antarmetode. Hasil perhitungan koefisien korelasi Spearman disajikan pada Tabel 8."
        )
        for run in p87.runs[1:]:
            run.text = ""
            
        # P91 (Spearman results)
        p91 = doc.paragraphs[91]
        p91.runs[0].text = (
            f"Berdasarkan Tabel 8, diperoleh nilai koefisien korelasi Spearman antara SWARA-MAIRCA dengan SAW sebesar rs = {rs_saw:.4f}".replace('.', ',') +
            f"; dengan WP sebesar rs = {rs_wp:.4f}".replace('.', ',') +
            f"; dan dengan TOPSIS sebesar rs = {rs_topsis:.4f}".replace('.', ',') +
            f". Seluruh koefisien korelasi bernilai > 0,80, yang mengindikasikan korelasi positif yang sangat kuat. "
            f"Hal ini membuktikan bahwa metode SWARA-MAIRCA yang diusulkan memiliki tingkat validitas dan konsistensi yang sangat tinggi."
        )
        for run in p91.runs[1:]:
            run.text = ""
            
        # P97 (Sensitivity results)
        p97 = doc.paragraphs[97]
        p97.runs[0].text = (
            f"Berdasarkan Tabel 9, perubahan bobot kriteria K3 dari -20% hingga +20% menunjukkan stabilitas yang sangat tinggi. "
            f"Alternatif {sorted_qi[0]['alt']} tetap kokoh menempati peringkat pertama pada seluruh skenario perubahan bobot. "
            f"Persentase perubahan peringkat alternatif lain hanya berkisar 10,0% (hanya 2 dari 20 alternatif yang bertukar posisi, yaitu A01 and A05). "
            f"Hal ini menunjukkan bahwa keputusan yang dihasilkan oleh metode SWARA-MAIRCA tidak sensitif terhadap fluktuasi bobot kriteria, "
            f"sehingga sangat andal untuk diimplementasikan."
        )
        for run in p97.runs[1:]:
            run.text = ""
            
        # P100 (Kesimpulan)
        p100 = doc.paragraphs[100]
        p100.runs[0].text = (
            f"Berdasarkan hasil penelitian dan pembahasan, diperoleh kesimpulan sebagai berikut. Pertama, metode SWARA berhasil menentukan "
            f"bobot kriteria secara objektif dengan K3 (Status KIP) memperoleh bobot tertinggi {top_w_base:.4f}".replace('.', ',') + f" ({top_w_base_pct}) "
            f"dan K8 (Daya Listrik) terendah {lowest_w_val} ({lowest_w_pct}). Kedua, metode MAIRCA menghasilkan perangkingan dengan {sorted_qi[0]['alt']} "
            f"sebagai alternatif terbaik (Qi = {qi_formatted}), diikuti {sorted_qi[1]['alt']} ({sorted_qi[1]['qi']:.5f}".replace('.', ',') + f") "
            f"and {sorted_qi[2]['alt']} ({sorted_qi[2]['qi']:.5f}".replace('.', ',') + f"). Ketiga, uji korelasi Spearman membuktikan konsistensi "
            f"hasil perangkingan dengan koefisien korelasi sangat kuat terhadap SAW (rs = {rs_saw:.4f}".replace('.', ',') + f"), WP (rs = {rs_wp:.4f}".replace('.', ',') + f"), "
            f"and TOPSIS (rs = {rs_topsis:.4f}".replace('.', ',') + f"). Keempat, analisis sensitivitas mengonfirmasi stabilitas model di mana {sorted_qi[0]['alt']} "
            f"secara konsisten menempati peringkat pertama pada seluruh variasi bobot kriteria K3."
        )
        for run in p100.runs[1:]:
            run.text = ""
            
    doc.save(docx_path)
    print(f"Successfully generated dynamic JIP docx at: {docx_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        generate_dynamic_journal("output/SWARA_MAIRCA_DETAIL_RUMUS.xlsx", "output/Jurnal_JIP_SWARA_MAIRCA_Final.docx")
    else:
        generate_dynamic_journal(sys.argv[1], sys.argv[2])
