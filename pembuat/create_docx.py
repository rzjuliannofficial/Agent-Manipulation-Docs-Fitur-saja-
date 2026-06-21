import argparse
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    """Sets margins (padding) inside a table cell in DXA units (20 DXA = 1 Pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    """Fills a cell's background with a hex color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_page_number(run):
    """Injects a dynamic Page field code into a text run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def create_premium_docx(filename="corporate_report.docx"):
    """Generates a premium Word document that looks crafted by a professional typographer/designer."""
    try:
        doc = Document()
        
        # Page Setup (Letter, 1 inch margins)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
            # Setup headers/footers
            footer = section.footer
            f_p = footer.paragraphs[0]
            f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            f_run = f_p.add_run("Internal Use Only | Page ")
            f_run.font.name = 'Arial'
            f_run.font.size = Pt(9)
            f_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
            add_page_number(f_p.add_run())

        # Base Typography Styles
        style_normal = doc.styles['Normal']
        style_normal.font.name = 'Arial'
        style_normal.font.size = Pt(10.5)
        style_normal.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)  # Cool grey-black
        
        # --- TITLE & SUBTITLE (Modern Minimalist Front Block) ---
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_before = Pt(36)
        title_p.paragraph_format.space_after = Pt(4)
        title_run = title_p.add_run("Q1 FINANCIAL & ENGINEERING REPORT")
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(26)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)  # Classic Navy
        
        sub_p = doc.add_paragraph()
        sub_p.paragraph_format.space_after = Pt(24)
        sub_run = sub_p.add_run("Strategic analysis of document pipelines and engineering milestones.")
        sub_run.font.name = 'Arial'
        sub_run.font.size = Pt(12)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        
        # Thin Divider Rule
        div_p = doc.add_paragraph()
        div_p.paragraph_format.space_after = Pt(36)
        div_p_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="1B365D"/></w:pBdr>')
        div_p._p.get_or_add_pPr().append(div_p_border)

        # --- SECTION 1: EXECUTIVE SUMMARY ---
        h1 = doc.add_heading(level=1)
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(6)
        h1_run = h1.add_run("1. Executive Summary")
        h1_run.font.name = 'Arial'
        h1_run.font.size = Pt(16)
        h1_run.font.bold = True
        h1_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        p1 = doc.add_paragraph(
            "This document presents the technical analysis and financial metrics of Q1 deliverables. "
            "Our team successfully deployed core scripting pipelines that read, parse, and generate "
            "formatted outputs for Microsoft Office standards and Adobe PDF layouts."
        )
        p1.paragraph_format.line_spacing = 1.15
        p1.paragraph_format.space_after = Pt(12)

        # --- SECTION 2: ACHIEVEMENTS ---
        h2 = doc.add_heading(level=2)
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(4)
        h2_run = h2.add_run("1.1 Core Engineering Milestones")
        h2_run.font.name = 'Arial'
        h2_run.font.size = Pt(13)
        h2_run.font.bold = True
        h2_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        bullets = [
            "Completed automated pipeline for parsing high-complexity Adobe PDF structures.",
            "Created openpyxl scripting classes to output premium corporate-standard spreadsheets.",
            "Integrated python-docx template handlers for report generation without AI visual footprint."
        ]
        for bullet in bullets:
            bp = doc.add_paragraph(bullet, style='List Bullet')
            bp.paragraph_format.space_after = Pt(4)
            bp.paragraph_format.line_spacing = 1.15
            
        doc.add_paragraph().paragraph_format.space_after = Pt(12) # Gap

        # --- SECTION 3: DELIVERABLES TABLE ---
        h1_2 = doc.add_heading(level=1)
        h1_2.paragraph_format.space_before = Pt(18)
        h1_2.paragraph_format.space_after = Pt(8)
        h1_2_run = h1_2.add_run("2. Project Status Matrix")
        h1_2_run.font.name = 'Arial'
        h1_2_run.font.size = Pt(16)
        h1_2_run.font.bold = True
        h1_2_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

        table_data = [
            ["Deliverable / Module", "Technology Stack", "Status"],
            ["PDF Analysis Module", "PyPDF & Python Core", "Completed"],
            ["Excel Engine", "Openpyxl / Data Science", "Completed"],
            ["Word Layout Engine", "Python-docx / Styling", "Completed"]
        ]
        
        table = doc.add_table(rows=len(table_data), cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Normal Table'
        
        # Set Table widths & borders
        tblPr = table._tbl.tblPr
        tblBorders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="BDC3C7"/>'
            f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="1B365D"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E8E8"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(tblBorders)

        for r_idx, row in enumerate(table.rows):
            trPr = row._tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '400') # Row height in dxa
            trPr.append(trHeight)
            
            for c_idx, val in enumerate(table_data[r_idx]):
                cell = row.cells[c_idx]
                cell.text = val
                cell.width = Inches(2.1)
                
                # Padding inside cells (120 dxa top/bottom, 150 dxa left/right)
                set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                
                # Format Paragraphs inside cells
                cell_p = cell.paragraphs[0]
                cell_p.paragraph_format.space_before = Pt(0)
                cell_p.paragraph_format.space_after = Pt(0)
                
                # Header row formatting
                if r_idx == 0:
                    set_cell_shading(cell, "1B365D")  # Premium navy header background
                    for run in cell_p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                else:
                    # Zebra striping on odd rows
                    if r_idx % 2 == 1:
                        set_cell_shading(cell, "F8F9FA")
                    for run in cell_p.runs:
                        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        doc.save(filename)
        print(f"Premium Word document successfully created: {filename}")
    except Exception as e:
        print(f"Error creating premium Word document: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Create a premium Word document.")
    parser.add_argument("-o", "--output", default="corporate_report.docx", help="Filename of the document to generate.")
    args = parser.parse_args()
    create_premium_docx(args.output)
