import argparse
import sys
from docx import Document

def extract_text_from_docx(docx_path, output_path=None):
    """Extracts all text from a Word (.docx) file and prints it or writes to a file."""
    try:
        # Reconfigure standard output to prevent Windows console encoding errors
        if hasattr(sys.stdout, 'reconfigure'):
            sys.stdout.reconfigure(encoding='utf-8')
        if hasattr(sys.stderr, 'reconfigure'):
            sys.stderr.reconfigure(encoding='utf-8')
            
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # Extract text from tables inside the Word document
        for table in doc.tables:
            full_text.append("\n--- Table Data ---")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    full_text.append(" | ".join(row_text))
            full_text.append("------------------\n")
                    
        content = "\n".join(full_text)
        if output_path:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            print(f"Successfully extracted text to: {output_path}")
        else:
            print(content)
    except Exception as e:
        print(f"Error reading Word document: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract text from a Word (.docx) file.")
    parser.add_argument("docx_path", help="Path to the .docx file to read.")
    parser.add_argument("-o", "--output", help="Optional path to save the extracted text.")
    args = parser.parse_args()
    extract_text_from_docx(args.docx_path, args.output)
