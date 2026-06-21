import docx

def compare_docs(path1, path2):
    doc1 = docx.Document(path1)
    doc2 = docx.Document(path2)
    
    print(f"Doc 1: {path1}")
    print(f"Number of paragraphs: {len(doc1.paragraphs)}")
    print(f"Number of tables: {len(doc1.tables)}")
    print(f"Number of sections: {len(doc1.sections)}")
    for i, sec in enumerate(doc1.sections):
        print(f"  Sec {i}: columns = {len(sec._sectPr.xpath('.//w:col')) or 1}")
        
    print(f"\nDoc 2: {path2}")
    print(f"Number of paragraphs: {len(doc2.paragraphs)}")
    print(f"Number of tables: {len(doc2.tables)}")
    print(f"Number of sections: {len(doc2.sections)}")
    for i, sec in enumerate(doc2.sections):
        print(f"  Sec {i}: columns = {len(sec._sectPr.xpath('.//w:col')) or 1}")
        
    # Print the first 10 paragraphs of each
    print("\n--- Doc 1 First 10 Paragraphs ---")
    for i in range(min(15, len(doc1.paragraphs))):
        print(f"P{i}: {doc1.paragraphs[i].text[:80]}")
        
    print("\n--- Doc 2 First 10 Paragraphs ---")
    for i in range(min(15, len(doc2.paragraphs))):
        print(f"P{i}: {doc2.paragraphs[i].text[:80]}")

compare_docs("output/Jurnal_JIP_SWARA_MAIRCA_Final.docx", "output/Jurnal_JIP_SWARA_MAIRCA_Final_2Kolom.docx")
