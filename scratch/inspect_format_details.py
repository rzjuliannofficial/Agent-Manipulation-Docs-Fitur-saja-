import docx

def inspect_run_formatting():
    doc1 = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final.docx")
    doc2 = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final_2Kolom.docx")
    
    print(f"Doc1 paragraphs: {len(doc1.paragraphs)}, Doc2 paragraphs: {len(doc2.paragraphs)}")
    
    differences = 0
    for i in range(min(len(doc1.paragraphs), len(doc2.paragraphs))):
        p1 = doc1.paragraphs[i]
        p2 = doc2.paragraphs[i]
        
        # We only care about runs
        r1_list = []
        for r in p1.runs:
            r1_list.append({
                "text": r.text,
                "font_name": r.font.name,
                "font_size": r.font.size.pt if r.font.size else None,
                "bold": r.font.bold,
                "italic": r.font.italic
            })
            
        r2_list = []
        for r in p2.runs:
            r2_list.append({
                "text": r.text,
                "font_name": r.font.name,
                "font_size": r.font.size.pt if r.font.size else None,
                "bold": r.font.bold,
                "italic": r.font.italic
            })
            
        # Compare them
        if len(r1_list) != len(r2_list):
            print(f"P{i} has different number of runs: Doc1={len(r1_list)}, Doc2={len(r2_list)}")
            differences += 1
        else:
            for idx in range(len(r1_list)):
                run1 = r1_list[idx]
                run2 = r2_list[idx]
                # Compare font, size, bold, italic (not text)
                if (run1["font_name"] != run2["font_name"] or 
                    run1["font_size"] != run2["font_size"] or 
                    run1["bold"] != run2["bold"] or 
                    run1["italic"] != run2["italic"]):
                    print(f"P{i} run {idx} formatting difference:")
                    print(f"  Doc1: font={run1['font_name']}, size={run1['font_size']}, bold={run1['bold']}, italic={run1['italic']}, text='{run1['text'][:30]}'")
                    print(f"  Doc2: font={run2['font_name']}, size={run2['font_size']}, bold={run2['bold']}, italic={run2['italic']}, text='{run2['text'][:30]}'")
                    differences += 1
                    break
        if differences >= 20:
            print("Too many differences, stopping.")
            break

inspect_run_formatting()
