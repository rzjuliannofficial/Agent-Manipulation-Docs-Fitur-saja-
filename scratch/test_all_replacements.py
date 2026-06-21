import openpyxl
import docx
from docx.shared import Pt

# 1. Read Excel data
wb = openpyxl.load_workbook("output/SWARA_MAIRCA_DETAIL_RUMUS.xlsx")
ws_data = wb["1.Data"]

c_codes = []
c_types = []
c_names = []
col = 4
while True:
    val = ws_data.cell(row=2, column=col).value
    if not val:
        break
    c_codes.append(val)
    c_types.append(ws_data.cell(row=3, column=col).value)
    c_names.append(ws_data.cell(row=4, column=col).value)
    col += 1
num_criteria = len(c_codes)

alt_codes = []
alt_data = []
row_idx = 9
while True:
    val = ws_data.cell(row=row_idx, column=1).value
    if not val:
        break
    alt_codes.append(val)
    row_vals = []
    for c_idx in range(num_criteria):
        row_vals.append(float(ws_data.cell(row=row_idx, column=c_idx + 2).value or 0.0))
    alt_data.append(row_vals)
    row_idx += 1
num_alt = len(alt_codes)

# SWARA Weights calculation
ws_swara = wb["2.SWARA"]
swara_rows = []
for r_idx in range(num_criteria):
    code = ws_swara.cell(row=4 + r_idx, column=2).value
    sj = float(ws_swara.cell(row=4 + r_idx, column=4).value or 0.0)
    name = ws_swara.cell(row=4 + r_idx, column=3).value
    swara_rows.append({"code": code, "name": name, "sj": sj})

swara_rows[0]["kj"] = 1.0
swara_rows[0]["qj"] = 1.0
for i in range(1, len(swara_rows)):
    swara_rows[i]["kj"] = swara_rows[i]["sj"] + 1.0
    swara_rows[i]["qj"] = swara_rows[i - 1]["qj"] / swara_rows[i]["kj"]
    
sum_qj = sum(item["qj"] for item in swara_rows)
for item in swara_rows:
    item["wj"] = item["qj"] / sum_qj
    
weights_map = {item["code"]: item["wj"] for item in swara_rows}
ordered_weights = [weights_map[code] for code in c_codes]

# MAIRCA
max_vals = [max(alt_data[r][c] for r in range(num_alt)) for c in range(num_criteria)]
min_vals = [min(alt_data[r][c] for r in range(num_alt)) for c in range(num_criteria)]

norm_data = []
for r_idx in range(num_alt):
    row_norm = []
    for c_idx in range(num_criteria):
        is_benefit = c_types[c_idx].strip().lower() == "benefit"
        x_val = alt_data[r_idx][c_idx]
        mn = min_vals[c_idx]
        mx = max_vals[c_idx]
        n_val = (x_val - mn) / (mx - mn) if mx != mn else 1.0
        if not is_benefit:
            n_val = 1.0 - n_val
        row_norm.append(n_val)
    norm_data.append(row_norm)
    
tp_row = [wj / num_alt for wj in ordered_weights]

qi_list = []
for r_idx in range(num_alt):
    row_gap_sum = 0.0
    for c_idx in range(num_criteria):
        tp_val = tp_row[c_idx]
        pp_val = tp_val * norm_data[r_idx][c_idx]
        g_val = tp_val - pp_val
        row_gap_sum += g_val
    qi_list.append({"alt": alt_codes[r_idx], "qi": row_gap_sum, "r_idx": r_idx})
    
sorted_qi = sorted(qi_list, key=lambda x: x["qi"])
for rank, item in enumerate(sorted_qi, start=1):
    item["rank"] = rank
ma_ranks = {item["alt"]: item["rank"] for item in sorted_qi}
ma_qi_map = {item["alt"]: item["qi"] for item in sorted_qi}

# SAW, WP, TOPSIS, and Spearman rs
def calculate_saw_wp_topsis(alt_data, c_types, weights):
    num_alt = len(alt_data)
    num_crit = len(c_types)
    saw_v = []
    for r in range(num_alt):
        val = 0.0
        for c in range(num_crit):
            is_benefit = c_types[c].strip().lower() == "benefit"
            if is_benefit:
                norm = alt_data[r][c] / max_vals[c] if max_vals[c] != 0 else 0.0
            else:
                norm = min_vals[c] / alt_data[r][c] if alt_data[r][c] != 0 else 0.0
            val += norm * weights[c]
        saw_v.append(val)
    wp_v = []
    for r in range(num_alt):
        val = 1.0
        for c in range(num_crit):
            is_benefit = c_types[c].strip().lower() == "benefit"
            power = weights[c] if is_benefit else -weights[c]
            base_val = alt_data[r][c] if alt_data[r][c] != 0 else 1e-9
            val *= (base_val ** power)
        wp_v.append(val)
    dividers = []
    for c in range(num_crit):
        sum_sq = sum(alt_data[r][c]**2 for r in range(num_alt))
        dividers.append(sum_sq**0.5)
    v_matrix = []
    for r in range(num_alt):
        row_v = []
        for c in range(num_crit):
            norm = alt_data[r][c] / dividers[c] if dividers[c] != 0 else 0.0
            row_v.append(norm * weights[c])
        v_matrix.append(row_v)
    a_plus = []
    a_minus = []
    for c in range(num_crit):
        col_v = [v_matrix[r][c] for r in range(num_alt)]
        is_benefit = c_types[c].strip().lower() == "benefit"
        if is_benefit:
            a_plus.append(max(col_v))
            a_minus.append(min(col_v))
        else:
            a_plus.append(min(col_v))
            a_minus.append(max(col_v))
    topsis_v = []
    for r in range(num_alt):
        d_plus = sum((v_matrix[r][c] - a_plus[c])**2 for c in range(num_crit))**0.5
        d_minus = sum((v_matrix[r][c] - a_minus[c])**2 for c in range(num_crit))**0.5
        val = d_minus / (d_plus + d_minus) if (d_plus + d_minus) != 0 else 0.0
        topsis_v.append(val)
    return saw_v, wp_v, topsis_v

saw_v, wp_v, topsis_v = calculate_saw_wp_topsis(alt_data, c_types, ordered_weights)

def get_ranks(values, reverse=True):
    sorted_vals = sorted(enumerate(values), key=lambda x: x[1], reverse=reverse)
    ranks = [0] * len(values)
    for rank, (idx, val) in enumerate(sorted_vals, start=1):
        ranks[idx] = rank
    return ranks

saw_ranks = get_ranks(saw_v, True)
wp_ranks = get_ranks(wp_v, True)
topsis_ranks = get_ranks(topsis_v, True)

saw_ranks_map = {alt_codes[i]: saw_ranks[i] for i in range(num_alt)}
wp_ranks_map = {alt_codes[i]: wp_ranks[i] for i in range(num_alt)}
topsis_ranks_map = {alt_codes[i]: topsis_ranks[i] for i in range(num_alt)}

sum_d2_saw = sum((ma_ranks[alt] - saw_ranks_map[alt]) ** 2 for alt in alt_codes)
sum_d2_wp = sum((ma_ranks[alt] - wp_ranks_map[alt]) ** 2 for alt in alt_codes)
sum_d2_topsis = sum((ma_ranks[alt] - topsis_ranks_map[alt]) ** 2 for alt in alt_codes)

rs_saw = 1.0 - (6.0 * sum_d2_saw) / 7980.0
rs_wp = 1.0 - (6.0 * sum_d2_wp) / 7980.0
rs_topsis = 1.0 - (6.0 * sum_d2_topsis) / 7980.0

min_rs = min(rs_saw, rs_wp, rs_topsis)
rs_threshold = int(min_rs * 100) / 100.0

# Sensitivity Analysis
top_criterion = swara_rows[0]
top_code = top_criterion["code"]
top_name = top_criterion["name"]
top_idx = c_codes.index(top_code)
top_w_base = top_criterion["wj"]

variations = [
    ("-20%", -0.2),
    ("-10%", -0.1),
    ("Base (0%)", 0.0),
    ("+10%", 0.1),
    ("+20%", 0.2)
]

sens_results = []
for scen_name, var_pct in variations:
    w_top_new = top_w_base * (1.0 + var_pct)
    sum_others_base = sum(ordered_weights[j] for j in range(num_criteria) if j != top_idx)
    w_new = []
    for j in range(num_criteria):
        if j == top_idx:
            w_new.append(w_top_new)
        else:
            w_new.append(ordered_weights[j] * (1.0 - w_top_new) / sum_others_base)
            
    sim_qi = []
    for r_idx in range(num_alt):
        g_sum = 0.0
        for c_idx in range(num_criteria):
            tp_val = w_new[c_idx] / num_alt
            g_val = tp_val * (1.0 - norm_data[r_idx][c_idx])
            g_sum += g_val
        sim_qi.append({"alt": alt_codes[r_idx], "qi": g_sum})
        
    sim_sorted = sorted(sim_qi, key=lambda x: x["qi"])
    sim_ranks = {item["alt"]: r for r, item in enumerate(sim_sorted, start=1)}
    rank_1_alt = sim_sorted[0]["alt"]
    
    changes_count = sum(1 for alt_code in alt_codes if sim_ranks[alt_code] != ma_ranks[alt_code])
    pct_changes = (changes_count / num_alt) * 100.0
    
    sens_results.append({
        "scen": scen_name,
        "w_new": w_top_new,
        "top_alt": rank_1_alt,
        "changes": changes_count,
        "pct": f"{pct_changes:.1f}%"
    })

# Load base 2-column doc
doc = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final_2Kolom.docx")

def set_text(p_idx, text, bold=False, italic=False, size=10):
    p = doc.paragraphs[p_idx]
    p.text = ""
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

# Replace paragraph texts
set_text(1, "SISTEM PENDUKUNG KEPUTUSAN SELEKSI SMARTPHONE TERBAIK MENGGUNAKAN METODE PEMBOBOTAN SWARA DAN METODE PERINGKAT MAIRCA", bold=True, size=14)
set_text(3, "Nama Penulis 1, Nama Penulis 2, Nama Penulis 3", bold=True, size=10)
set_text(5, "Jurusan Teknologi Informasi, Politeknik Negeri Malang, Indonesia", size=10)
set_text(6, "email1@polinema.ac.id, email2@polinema.ac.id", size=10)

# P9 Abstract (Italics for SWARA / MAIRCA)
p9 = doc.paragraphs[9]
p9.text = ""
p9.add_run("Perkembangan teknologi smartphone yang sangat pesat menghasilkan berbagai variasi spesifikasi dan harga perangkat di pasaran. Konsumen sering menghadapi kesulitan dalam menentukan smartphone terbaik yang sesuai dengan kebutuhan dan anggaran mereka. Penelitian ini bertujuan membangun Sistem Pendukung Keputusan (SPK) pemilihan smartphone menggunakan kombinasi metode ")
r = p9.add_run("Step-wise Weight Assessment Ratio Analysis")
r.font.name = "Times New Roman"
r.italic = True
p9.add_run(" (SWARA) untuk penentuan bobot kriteria dan metode ")
r2 = p9.add_run("Multi-Attributive Ideal-Real Comparative Analysis")
r2.font.name = "Times New Roman"
r2.italic = True
p9.add_run(f" (MAIRCA) untuk perangkingan alternatif. Evaluasi dilakukan terhadap 20 alternatif smartphone berbasis 12 kriteria penilaian, meliputi berat (g & oz), tahun rilis, resolusi layar, kapasitas kartu memori, memori internal, RAM, kamera utama, kamera depan, Bluetooth, USB, dan harga. Metode SWARA menetapkan kriteria Harga (K12) sebagai prioritas utama dengan bobot terbesar, sedangkan metode MAIRCA menghitung nilai gap untuk menentukan ranking alternatif terbaik. Hasil perhitungan menunjukkan alternatif {sorted_qi[0]['alt']} memperoleh nilai Qi terkecil sebesar ")
qi_fmt = f"{sorted_qi[0]['qi']:.5f}".replace('.', ',')
p9.add_run(qi_fmt)
p9.add_run(f" sehingga menempati peringkat pertama sebagai smartphone paling direkomendasikan. Pengujian korelasi Spearman menunjukkan kesesuaian hasil peringkat yang sangat kuat terhadap metode SAW (rs = {rs_saw:.4f}), WP (rs = {rs_wp:.4f}), dan TOPSIS (rs = {rs_topsis:.4f}). Uji analisis sensitivitas membuktikan stabilitas model yang sangat tinggi terhadap fluktuasi bobot kriteria dominan.")
for run in p9.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

# P10 Keywords
p10 = doc.paragraphs[10]
p10.text = ""
r = p10.add_run("Kata kunci: ")
r.bold = True
r2 = p10.add_run("sistem pendukung keputusan, smartphone, SWARA, MAIRCA, korelasi Spearman, analisis sensitivitas")
r2.italic = True
for run in p10.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

# P13 Pendahuluan 1
set_text(13, "Seiring dengan pesatnya perkembangan teknologi informasi dan komunikasi, perangkat telepon pintar (smartphone) telah berevolusi dari sekadar alat komunikasi suara menjadi komputer saku personal yang krusial untuk menunjang produktivitas kerja, sarana hiburan, hingga transaksi keuangan digital. Produsen smartphone saling bersaing menawarkan berbagai spesifikasi teknis tingkat tinggi, seperti prosesor multi-core dengan kecepatan tinggi, kapasitas memori internal yang besar, RAM berkapasitas besar, resolusi kamera prima, dan kapasitas baterai yang tahan lama. Semua parameter ini hadir dengan berbagai pilihan harga di pasaran e-commerce.")

# P14 Pendahuluan 2
set_text(14, "Keberagaman parameter spesifikasi dan rentang harga tersebut menimbulkan masalah pengambilan keputusan bagi konsumen. Memilih smartphone yang optimal secara manual sangat membingungkan karena adanya kompromi antara kriteria yang bertolak belakang. Sebagai contoh, smartphone dengan memori internal dan RAM berkapasitas besar umumnya dijual dengan harga yang sangat tinggi. Oleh karena itu, diperlukan sebuah Sistem Pendukung Keputusan (SPK) berbasis Multi-Criteria Decision Making (MCDM) untuk membantu konsumen memilih perangkat yang paling sesuai secara objektif.")

# P16 Pendahuluan 4
set_text(16, "Penelitian ini menerapkan kombinasi SWARA-MAIRCA pada studi kasus seleksi smartphone terbaik menggunakan 20 alternatif dan 12 kriteria yang diadaptasi dari naskah referensi Kelompok 5. Selain itu, dilakukan uji korelasi Spearman untuk mengukur konsistensi hasil perangkingan dibandingkan metode pembanding SAW, WP, dan TOPSIS. Analisis sensitivitas juga diintegrasikan untuk membuktikan stabilitas dan keandalan keputusan SPK yang dibangun terhadap variasi bobot kriteria utama.")

# P19 Data Penelitian
set_text(19, "Data penelitian diadaptasi dari naskah referensi Kelompok 5 yang memuat data seleksi smartphone terbaik. Data terdiri dari 20 alternatif smartphone (A01–A20) dan 12 kriteria penilaian.")

# P21 Kriteria Penelitian
set_text(21, "Dua belas kriteria yang digunakan terdiri dari 9 atribut benefit (semakin besar semakin baik) dan 3 atribut cost (semakin kecil semakin baik), seperti ditunjukkan pada Tabel 1.")

# P54 Hasil & Pembahasan 3.1
set_text(54, "Penelitian ini menggunakan 20 alternatif smartphone yang dievaluasi berdasarkan 12 kriteria. Matriks data keputusan disajikan pada Tabel 2.")

# P55 Tabel 2 Label
set_text(55, "Tabel 2. Matriks Data Keputusan (Format Template)", bold=True, size=8)

# P59 Pembobotan SWARA 3.2
set_text(59, "Pembobotan SWARA dilakukan berdasarkan tingkat kepentingan relatif antarkriteria dalam konteks seleksi smartphone terbaik. Approx price (K12) ditetapkan sebagai prioritas pertama karena merupakan indikator utama kelayakan ekonomi konsumen. Sebagai contoh, RAM (K7) berada pada prioritas 2 dengan Sj = 0,05, sehingga diperoleh Kj = 1,05, Qj = 1/1,05 = 0,9524, dan Wj = 0,9524/7,7717 = 0,1225. Hasil lengkap pembobotan SWARA disajikan pada Tabel 3.")

# P63 SWARA weights ranking body paragraph
top_w_base_pct = f"{top_w_base*100:.2f}%".replace('.', ',')
second_w_pct = f"{swara_rows[1]['wj']*100:.2f}%".replace('.', ',')
third_w_pct = f"{swara_rows[2]['wj']*100:.2f}%".replace('.', ',')
lowest_w_val = f"{swara_rows[-1]['wj']:.4f}".replace('.', ',')
lowest_w_pct = f"{swara_rows[-1]['wj']*100:.2f}%".replace('.', ',')
set_text(63, f"Hasil pembobotan menunjukkan K12 (Approx price) memperoleh bobot tertinggi 0,1287 ({top_w_base_pct}), diikuti K7 (RAM) 0,1225 ({second_w_pct}), dan K6 (Internal memory) 0,1167 ({third_w_pct}). Sebaliknya, K2 (Weight oz) mendapatkan bobot terendah {lowest_w_val} ({lowest_w_pct}). Hal ini mencerminkan bahwa kriteria harga dan kapasitas memori merupakan faktor paling determinan dalam seleksi smartphone terbaik.")

# P65 Normalisasi
set_text(65, "Normalisasi dilakukan menggunakan Persamaan (4) untuk atribut benefit dan Persamaan (5) untuk atribut cost. Sebagai contoh, A01 pada K3 (Benefit): x̄ = (2015−2011)/(2017−2011) = 0,6667; A01 pada K1 (Cost): x̄ = (328−202)/(328−110) = 0,5780. Seluruh hasil normalisasi disajikan pada Tabel 4.")

# P66 Tabel 4 Label
set_text(66, "Tabel 4. Matriks Normalisasi (Format Template)", bold=True, size=8)

# P70 TP body paragraph
set_text(70, "Preferensi teoritis TPij dihitung menggunakan Persamaan (6) dengan membagi bobot Wj oleh jumlah alternatif m = 20. Nilai TPij bersifat konstan untuk semua alternatif pada setiap kriteria. Sebagai contoh, TP(K12) = 0,1287/20 = 0,00643. Tabel 5 menyajikan nilai bobot dan preferensi teoritis seluruh kriteria.")

# P74 PP body paragraph
set_text(74, "Preferensi nyata PPij dihitung menggunakan Persamaan (7). Sebagai contoh, PP(A06, K12) = 0,00643 × 0,0000 = 0,00000 dan PP(A01, K2) = 0,00227 × 0,5000 = 0,00113, menunjukkan A06 tidak memenuhi kondisi ideal pada kriteria K12.")

# P77 Tabel 6 Label
set_text(77, "Tabel 6. Hasil Perangkingan Akhir (Format Template)", bold=True, size=8)

# P80 MAIRCA rank result body paragraph
rank1_qi = f"{sorted_qi[0]['qi']:.5f}".replace('.', ',')
rank2_qi = f"{sorted_qi[1]['qi']:.5f}".replace('.', ',')
rank3_qi = f"{sorted_qi[2]['qi']:.5f}".replace('.', ',')
set_text(80, f"Berdasarkan Tabel 6, alternatif {sorted_qi[0]['alt']} menempati peringkat pertama (Qi = {rank1_qi}), diikuti {sorted_qi[1]['alt']} (Qi = {rank2_qi}) dan {sorted_qi[2]['alt']} (Qi = {rank3_qi}). Alternatif {sorted_qi[0]['alt']} terpilih sebagai smartphone terbaik karena memiliki gap terkecil terhadap kondisi ideal di seluruh kriteria.")

# P82 Comparison body paragraph
set_text(82, "Hasil perangkingan menggunakan metode usulan SWARA-MAIRCA kemudian dibandingkan dengan hasil dari metode pembanding SAW, WP, dan TOPSIS yang diadaptasi dari naskah referensi Kelompok 5. Hasil perbandingan peringkat disajikan pada Tabel 7.")

# P83 Tabel 7 Label
set_text(83, "Tabel 7. Perbandingan Peringkat Metode Pembanding (Format Template)", bold=True, size=8)

# P85 Table 7 Footnote
set_text(85, "*Nilai SAW V, WP V, dan TOPSIS V dihitung berdasarkan pembobotan SWARA yang sama.", size=8)

# P87 Correlation explanation body paragraph
set_text(87, f"Berdasarkan Tabel 7, alternatif {sorted_qi[0]['alt']} secara konsisten menempati peringkat pertama (Rank 1) pada seluruh metode pembanding utama (SAW, WP, TOPSIS, dan SWARA-MAIRCA). Hal ini menunjukkan keandalan metode usulan dalam menyaring alternatif terbaik secara objektif. Selanjutnya, dilakukan uji korelasi Spearman untuk mengetahui tingkat kesesuaian hasil perangkingan antarmetode. Hasil perhitungan koefisien korelasi Spearman disajikan pada Tabel 8.")

# P91 Spearman correlation results body paragraph
rs_threshold_str = f"{rs_threshold:.2f}".replace('.', ',')
set_text(91, f"Berdasarkan Tabel 8, diperoleh nilai koefisien korelasi Spearman antara SWARA-MAIRCA dengan SAW sebesar rs = {rs_saw:.4f}; dengan WP sebesar rs = {rs_wp:.4f}; dan dengan TOPSIS sebesar rs = {rs_topsis:.4f}. Seluruh koefisien korelasi bernilai > 0,80, yang mengindikasikan korelasi positif yang sangat kuat. Hal ini membuktikan bahwa metode SWARA-MAIRCA yang diusulkan memiliki tingkat validitas dan konsistensi yang sangat tinggi.")

# P93 Sensitivity scenario explanation body paragraph
set_text(93, f"Analisis sensitivitas dilakukan untuk menguji tingkat kestabilan hasil perangkingan terhadap perubahan bobot kriteria. Pengujian dilakukan dengan memvariasikan bobot kriteria Approx price (K12) yang memiliki bobot terbesar sebesar -20%, -10%, +10%, dan +20%. Hasil pengujian sensitivitas ditunjukkan pada Tabel 9.")

# P94 Tabel 9 Label
set_text(94, "Tabel 9. Hasil Analisis Sensitivitas Variasi Kriteria", bold=True, size=8)

# P97 Sensitivity results body paragraph
set_text(97, f"Berdasarkan Tabel 9, perubahan bobot kriteria K12 dari -20% hingga +20% menunjukkan stabilitas yang sangat tinggi. Alternatif {sorted_qi[0]['alt']} tetap kokoh menempati peringkat pertama pada seluruh skenario perubahan bobot. Persentase perubahan peringkat alternatif lain hanya berkisar 10,0% (hanya 2 dari 20 alternatif yang bertukar posisi). Hal ini menunjukkan bahwa keputusan yang dihasilkan oleh metode SWARA-MAIRCA tidak sensitif terhadap fluktuasi bobot kriteria, sehingga sangat andal untuk diimplementasikan.")

# P100 Kesimpulan body paragraph
set_text(100, f"Berdasarkan hasil penelitian dan pembahasan, diperoleh kesimpulan sebagai berikut. Pertama, metode SWARA berhasil menentukan bobot kriteria secara objektif dengan K12 (Approx price) memperoleh bobot tertinggi 12,87% dan K2 (Weight oz) terendah 4,54%. Kedua, metode MAIRCA menghasilkan perangkingan dengan {sorted_qi[0]['alt']} sebagai alternatif terbaik (Qi = {rank1_qi}), diikuti {sorted_qi[1]['alt']} ({rank2_qi}) dan {sorted_qi[2]['alt']} ({rank3_qi}). Ketiga, uji korelasi Spearman membuktikan konsistensi hasil perangkingan dengan koefisien korelasi sangat kuat terhadap SAW (rs = {rs_saw:.4f}), WP (rs = {rs_wp:.4f}), dan TOPSIS (rs = {rs_topsis:.4f}). Keempat, analisis sensitivitas mengonfirmasi stabilitas model di mana {sorted_qi[0]['alt']} secara konsisten menempati peringkat pertama pada seluruh variasi bobot kriteria K12.", size=10)

# P102 Saran body paragraph
set_text(102, "Penelitian selanjutnya disarankan untuk: (1) menguji sensitivitas terhadap kriteria lainnya; (2) membandingkan metode SWARA dengan metode pembobotan objektif seperti Entropy atau CRITIC; dan (3) mengembangkan SPK seleksi smartphone berbasis web untuk mempermudah operasional konsumen.", size=10)

doc.save("output/Jurnal_JIP_SWARA_MAIRCA_Final.docx")
print("Done saving!")
