import os
import docx

for f in os.listdir("output"):
    if f.endswith(".docx"):
        path = os.path.join("output", f)
        try:
            doc = docx.Document(path)
            title = ""
            for p in doc.paragraphs:
                if p.text.strip():
                    title = p.text.strip()
                    break
            print(f"File: {f} | Title: {title} | Paragraphs: {len(doc.paragraphs)}")
        except Exception as e:
            print(f"File: {f} | Error: {e}")
