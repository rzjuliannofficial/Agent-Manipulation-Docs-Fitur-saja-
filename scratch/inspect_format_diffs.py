import docx

doc1 = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final.docx")
doc2 = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final_2Kolom.docx")

def format_sec(sec):
    sectPr = sec._sectPr
    cols = sectPr.xpath('.//w:cols')
    num_cols = cols[0].get(docx.oxml.ns.qn('w:num')) if cols else '1'
    space_cols = cols[0].get(docx.oxml.ns.qn('w:space')) if cols else '0'
    return {
        "w": sec.page_width.inches if sec.page_width else None,
        "h": sec.page_height.inches if sec.page_height else None,
        "top": sec.top_margin.inches if sec.top_margin else None,
        "bottom": sec.bottom_margin.inches if sec.bottom_margin else None,
        "left": sec.left_margin.inches if sec.left_margin else None,
        "right": sec.right_margin.inches if sec.right_margin else None,
        "cols": num_cols,
        "col_space": space_cols
    }

print("=== Section 0 Formatting ===")
print("Doc1:", format_sec(doc1.sections[0]))
print("Doc2:", format_sec(doc2.sections[0]))

print("\n=== Section 1 Formatting ===")
print("Doc1:", format_sec(doc1.sections[1]))
print("Doc2:", format_sec(doc2.sections[1]))

print("\n=== First 10 Paragraphs Details ===")
for i in range(15):
    p1 = doc1.paragraphs[i]
    p2 = doc2.paragraphs[i]
    
    def p_info(p):
        font_name = None
        font_size = None
        bold = None
        italic = None
        if p.runs:
            font_name = p.runs[0].font.name
            font_size = p.runs[0].font.size.pt if p.runs[0].font.size else None
            bold = p.runs[0].font.bold
            italic = p.runs[0].font.italic
        return {
            "text_len": len(p.text),
            "align": str(p.alignment),
            "before": p.paragraph_format.space_before.pt if p.paragraph_format.space_before else None,
            "after": p.paragraph_format.space_after.pt if p.paragraph_format.space_after else None,
            "line_spacing": p.paragraph_format.line_spacing,
            "font_name": font_name,
            "font_size": font_size,
            "bold": bold,
            "italic": italic
        }
        
    print(f"P{i} Doc1: {p_info(p1)}")
    print(f"P{i} Doc2: {p_info(p2)}")
