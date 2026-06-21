import docx

doc = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final_2Kolom.docx")

# Let's inspect P9 (Abstract)
p = doc.paragraphs[9]
print("P9 text:", p.text[:100])
print("Runs in P9:")
for r_idx, run in enumerate(p.runs):
    print(f"  Run {r_idx}: text='{run.text}' | bold={run.font.bold} | italic={run.font.italic}")
    
# Let's inspect P1 (Title)
p_title = doc.paragraphs[1]
print("\nP1 text:", p_title.text)
print("Runs in P1:")
for r_idx, run in enumerate(p_title.runs):
    print(f"  Run {r_idx}: text='{run.text}'")
