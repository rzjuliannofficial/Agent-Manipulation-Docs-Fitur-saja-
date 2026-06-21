import docx
import sys

def write_diffs():
    doc1 = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final.docx")
    doc2 = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final_2Kolom.docx")
    
    with open("scratch/diffs.txt", "w", encoding="utf-8") as f:
        f.write(f"Doc1 paragraphs: {len(doc1.paragraphs)}, Doc2 paragraphs: {len(doc2.paragraphs)}\n\n")
        
        for i in range(min(len(doc1.paragraphs), len(doc2.paragraphs))):
            t1 = doc1.paragraphs[i].text.strip()
            t2 = doc2.paragraphs[i].text.strip()
            if t1 != t2:
                f.write(f"--- Paragraph {i} ---\n")
                f.write(f"Final:     {t1}\n")
                f.write(f"2Kolom:    {t2}\n\n")

write_diffs()
