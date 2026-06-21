import docx

doc1 = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final.docx")
doc2 = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final_2Kolom.docx")

print("--- Doc1 Paragraphs ---")
for i in range(15):
    print(f"P{i}: {doc1.paragraphs[i].text[:100]}")
    
print("\n--- Doc2 Paragraphs ---")
for i in range(15):
    print(f"P{i}: {doc2.paragraphs[i].text[:100]}")
