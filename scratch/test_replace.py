import docx
from docx.shared import Pt

doc = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final_2Kolom.docx")

print(f"Total paragraphs in Doc2: {len(doc.paragraphs)}")

p_title = doc.paragraphs[1]
print(f"Original P1: '{p_title.text}'")

# Test modifying it
p_title.text = ""
run = p_title.add_run("SISTEM PENDUKUNG KEPUTUSAN SELEKSI SMARTPHONE TERBAIK MENGGUNAKAN METODE PEMBOBOTAN SWARA DAN METODE PERINGKAT MAIRCA")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.font.bold = True

print(f"Modified P1: '{doc.paragraphs[1].text}'")
print(f"Modified P1 alignment: {doc.paragraphs[1].alignment}")
print(f"Modified P1 before/after space: {doc.paragraphs[1].paragraph_format.space_before}, {doc.paragraphs[1].paragraph_format.space_after}")
