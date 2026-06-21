import docx
from docx.shared import Pt

def format_run(run, font_name="Times New Roman", size_pt=10, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    
    # XML font override
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(docx.oxml.ns.qn('w:ascii'), font_name)
    rFonts.set(docx.oxml.ns.qn('w:hAnsi'), font_name)

def test():
    doc = docx.Document()
    p = doc.add_paragraph()
    run = p.add_run("Hello Test")
    format_run(run, "Times New Roman", 14, True, False)
    
    doc.save("scratch/test_out.docx")
    
    # Read it back
    doc_in = docx.Document("scratch/test_out.docx")
    r = doc_in.paragraphs[0].runs[0]
    print("Font Name:", r.font.name)
    print("Font Size:", r.font.size.pt if r.font.size else None)
    print("Bold:", r.font.bold)
    print("Italic:", r.font.italic)

test()
