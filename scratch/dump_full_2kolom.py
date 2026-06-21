import docx

doc = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final_2Kolom.docx")
with open("scratch/full_2kolom_content.txt", "w", encoding="utf-8") as f:
    for idx, p in enumerate(doc.paragraphs):
        f.write(f"P{idx}: {p.text}\n")
print(f"Done. Wrote {len(doc.paragraphs)} paragraphs.")
