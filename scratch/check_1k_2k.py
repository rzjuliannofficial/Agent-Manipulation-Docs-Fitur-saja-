import docx

doc_1k = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final_1Kolom.docx")
doc_2k = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final_2Kolom.docx")

print(f"1Kolom length: {len(doc_1k.paragraphs)}")
print(f"2Kolom length: {len(doc_2k.paragraphs)}")

diffs = 0
for i in range(min(len(doc_1k.paragraphs), len(doc_2k.paragraphs))):
    t1 = doc_1k.paragraphs[i].text
    t2 = doc_2k.paragraphs[i].text
    if t1 != t2:
        print(f"Diff at P{i}:")
        print(f"  1Kolom: {t1[:100]}")
        print(f"  2Kolom: {t2[:100]}")
        diffs += 1
if diffs == 0:
    print("All paragraph text matches between 1Kolom and 2Kolom!")
