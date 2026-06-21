import docx

doc1 = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final.docx")
doc2 = docx.Document("output/Jurnal_JIP_SWARA_MAIRCA_Final_2Kolom.docx")

diffs = 0
for i in range(len(doc1.paragraphs)):
    t1 = doc1.paragraphs[i].text
    t2 = doc2.paragraphs[i].text
    if t1 != t2:
        print(f"Diff at P{i}:")
        print(f"  Doc1: {t1}")
        print(f"  Doc2: {t2}")
        diffs += 1

if diffs == 0:
    print("All paragraph texts match perfectly!")
